"""Document Parser Service - Parse PDF, DOCX, TXT, HTML, XML, images, audio, video into text chunks."""

import io
import logging
import re
import unicodedata
from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentParser:
    """Service for parsing various document formats into text."""

    SUPPORTED_TYPES = {
        # Text documents
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
        "text/markdown": "txt",
        "text/csv": "spreadsheet",
        "text/xml": "xml",
        "text/html": "html",
        "text/css": "txt",
        "text/javascript": "txt",
        "application/json": "txt",
        "application/xml": "xml",
        "application/x-sh": "txt",
        "application/x-www-form-urlencoded": "txt",
        "application/vnd.ms-excel": "spreadsheet",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "spreadsheet",
        # Images
        "image/png": "image",
        "image/jpeg": "image",
        "image/gif": "image",
        "image/webp": "image",
        "image/svg+xml": "image",
        # Audio
        "audio/mpeg": "audio",
        "audio/mp3": "audio",
        "audio/wav": "audio",
        "audio/x-wav": "audio",
        "audio/ogg": "audio",
        "audio/flac": "audio",
        "audio/x-flac": "audio",
        "audio/x-m4a": "audio",
        # Video
        "video/mp4": "video",
        "video/x-msvideo": "video",
        "video/avi": "video",
        "video/x-matroska": "video",
        "video/mkv": "video",
        "video/webm": "video",
    }

    # MIME types that should be skipped (not indexable)
    SKIP_TYPES = {
        "application/x-folder",
        "application/zip",
        "application/x-zip-compressed",
        "application/x-tar",
        "application/gzip",
        "application/x-rar-compressed",
        "application/x-7z-compressed",
        "application/octet-stream",
        "image/vnd.microsoft.icon",
        "application/x-executable",
        "application/x-sharedlib",
    }

    CHUNK_SIZE = 1000  # Characters per chunk
    CHUNK_OVERLAP = 200  # Overlap between chunks

    def get_file_category(self, mime_type: str) -> str | None:
        """Return file category: 'text', 'image', 'audio', 'video', or None."""
        doc_type = self.SUPPORTED_TYPES.get(mime_type)
        if doc_type in ("pdf", "docx", "txt", "html", "xml", "spreadsheet"):
            return "text"
        if doc_type == "image":
            return "image"
        if doc_type == "audio":
            return "audio"
        if doc_type == "video":
            return "video"
        return None

    def parse(self, content: bytes, content_type: str) -> str:
        """
        Parse document content to plain text.
        Only handles text documents (pdf/docx/txt/html/xml).
        For image/audio/video, use dedicated parsers.
        """
        doc_type = self.SUPPORTED_TYPES.get(content_type)

        if doc_type == "pdf":
            return self._parse_pdf(content)
        elif doc_type == "docx":
            return self._parse_docx(content)
        elif doc_type == "html":
            return self._parse_html(content)
        elif doc_type == "xml":
            return self._parse_xml(content)
        elif doc_type == "spreadsheet":
            return self._parse_spreadsheet(content, content_type)
        elif doc_type == "txt":
            return self._parse_txt(content)
        else:
            raise ValueError(f"Unsupported content type for text parsing: {content_type}")

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF with OCR fallback for garbled/empty pages."""
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        garbled_pages = []  # (page_index, ) for OCR retry

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip() and not self._is_garbled(text):
                text_parts.append((i, text))
            else:
                garbled_pages.append(i)

        # OCR fallback for garbled/empty pages (limit to first 20 pages to prevent memory explosion)
        MAX_OCR_PAGES = 20
        if garbled_pages:
            ocr_targets = garbled_pages[:MAX_OCR_PAGES]
            if len(garbled_pages) > MAX_OCR_PAGES:
                logger.warning("OCR limited to first %d of %d garbled pages", MAX_OCR_PAGES, len(garbled_pages))
            try:
                from pdf2image import convert_from_bytes
                import pytesseract

                for page_idx in ocr_targets:
                    try:
                        # Convert only the single page needed (not the entire PDF)
                        images = convert_from_bytes(
                            content, dpi=150, first_page=page_idx + 1, last_page=page_idx + 1,
                        )
                        if images:
                            ocr_text = pytesseract.image_to_string(images[0], lang="kor+eng")
                            if ocr_text and ocr_text.strip():
                                text_parts.append((page_idx, ocr_text))
                                logger.info(f"OCR fallback succeeded for page {page_idx + 1}")
                    except Exception as e:
                        logger.warning(f"OCR failed for page {page_idx + 1}: {e}")
            except ImportError:
                logger.warning("pdf2image/pytesseract not installed, skipping OCR fallback")
            except Exception as e:
                logger.warning(f"PDF to image conversion failed: {e}")

        # Sort by page index and join
        text_parts.sort(key=lambda x: x[0])
        return "\n\n".join(text for _, text in text_parts)

    def _is_garbled(self, text: str) -> bool:
        """Check if extracted text is garbled (unreadable unicode artifacts)."""
        if not text or not text.strip():
            return True

        normal_count = 0
        total_count = 0

        for char in text:
            if char.isspace():
                continue
            total_count += 1
            # ASCII printable
            if '\x20' <= char <= '\x7e':
                normal_count += 1
            # Korean (Hangul Syllables + Jamo + Compatibility Jamo)
            elif '\uac00' <= char <= '\ud7af' or '\u1100' <= char <= '\u11ff' or '\u3130' <= char <= '\u318f':
                normal_count += 1
            # CJK Unified Ideographs
            elif '\u4e00' <= char <= '\u9fff':
                normal_count += 1
            # Common punctuation and symbols
            elif unicodedata.category(char).startswith(('P', 'S', 'N')):
                normal_count += 1

        if total_count == 0:
            return True

        ratio = normal_count / total_count
        return ratio < 0.5

    def _parse_html(self, content: bytes) -> str:
        """Extract text from HTML, stripping tags and non-content elements."""
        from bs4 import BeautifulSoup

        text = self._parse_txt(content)
        soup = BeautifulSoup(text, "html.parser")

        # Remove non-content elements
        for tag in soup.find_all(["script", "style", "meta", "link", "noscript", "nav", "footer", "header"]):
            tag.decompose()

        extracted = soup.get_text(separator="\n", strip=True)

        # Collapse multiple blank lines
        extracted = re.sub(r"\n{3,}", "\n\n", extracted)
        return extracted

    def _parse_xml(self, content: bytes) -> str:
        """Extract text content from XML, stripping all tags."""
        from lxml import etree

        text = self._parse_txt(content)
        try:
            root = etree.fromstring(text.encode("utf-8"))
            # itertext() yields all text content in document order
            parts = [t.strip() for t in root.itertext() if t.strip()]
            return "\n".join(parts)
        except etree.XMLSyntaxError:
            # Fallback: regex strip tags
            cleaned = re.sub(r"<[^>]+>", " ", text)
            cleaned = re.sub(r"\s+", " ", cleaned).strip()
            return cleaned

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX."""
        doc = DocxDocument(io.BytesIO(content))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)

    def _parse_txt(self, content: bytes) -> str:
        """Decode plain text."""
        # Try common encodings
        for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        # Fallback with error handling
        return content.decode("utf-8", errors="replace")

    def chunk_text(
        self,
        text: str,
        chunk_size: int | None = None,
        overlap: int | None = None,
    ) -> list[str]:
        """Split text into overlapping chunks for embedding."""
        chunk_size = chunk_size or self.CHUNK_SIZE
        overlap = overlap or self.CHUNK_OVERLAP

        if len(text) <= chunk_size:
            return [text] if text.strip() else []

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                for sep in [". ", ".\n", "! ", "!\n", "? ", "?\n", "\n\n"]:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > chunk_size // 2:
                        end = start + last_sep + len(sep)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start position with overlap
            start = end - overlap
            if start >= len(text):
                break

        return chunks

    def _parse_spreadsheet(self, content: bytes, content_type: str) -> str:
        """Parse spreadsheet into plain text (fallback for non-chunked path)."""
        from app.services.document.spreadsheet_parser import SpreadsheetParser
        sp = SpreadsheetParser()
        chunks = sp.parse(content, content_type)
        return "\n\n".join(c.text for c in chunks)

    def parse_and_chunk(
        self,
        content: bytes,
        content_type: str,
    ) -> list[tuple[str, str | None]]:
        """Parse and chunk in one step. Returns [(chunk_text, section_name), ...].

        For spreadsheets, section_name is the sheet label (e.g. "Sheet: 매출현황").
        For other types, section_name is None.
        """
        doc_type = self.SUPPORTED_TYPES.get(content_type)

        if doc_type == "spreadsheet":
            from app.services.document.spreadsheet_parser import SpreadsheetParser
            sp = SpreadsheetParser()
            sheet_chunks = sp.parse(content, content_type)
            return [(c.text, f"Sheet: {c.sheet_name}") for c in sheet_chunks]

        # Non-spreadsheet: parse then chunk, section=None
        text = self.parse(content, content_type)
        chunks = self.chunk_text(text)
        return [(c, None) for c in chunks]

    def is_supported(self, content_type: str) -> bool:
        """Check if a content type is supported."""
        return content_type in self.SUPPORTED_TYPES


# Singleton
_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    """Get or create the document parser singleton."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
